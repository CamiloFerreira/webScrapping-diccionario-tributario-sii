from bs4 import BeautifulSoup 
import requests as rq
from docx import Document

url = "https://www.sii.cl/diccionario_tributario/dicc_a.htm"
letras = ["a","b","c","d","e","f","g","h","i","j","k","l","m","ñ","o","p","q","r","s","t","u","v","w","x","y","z"]
def ObtenerDatos(url,letra):
	response = rq.get(url)
	soup=BeautifulSoup(response.content,'html5lib')
	divs = soup.findAll("p")
	aLetra = []
	for div in divs : 
		aTitulo = div.find("b")
		if(aTitulo != None and aTitulo.text.lower() != letra):
			t_Titulo = "" 
			t_Titulo = aTitulo.text.strip()
			t_Espacio= t_Titulo.split("\n")
			if(len(t_Espacio) > 1):
				t_Final = ""
				for t in t_Espacio:
					t_Final +=t.strip()+" "			
				t_Titulo = t_Final			
			aTitulo.extract()
			contenido = div.text
			cEspacio = contenido.split("\n")
			contenido = ""
			for t in cEspacio:
				contenido += t.strip() +" " 	
			aLetra.append({"titulo":t_Titulo,"contenido":contenido})
	return aLetra
def generarDic():
	dic = {}
	for letra in letras:
		url = "https://www.sii.cl/diccionario_tributario/dicc_"+letra+".htm"
		if(len(ObtenerDatos(url,letra)) > 0 ):
			dic[letra] = ObtenerDatos(url,letra)
	return dic
#Crear documento
doc = Document()
dic = generarDic()
for letra in dic:
	doc.add_heading(letra.upper(), level=1)
	p = doc.add_paragraph()
	p.add_run("\n")
	for item in dic[letra]:
		p.add_run(item['titulo']).bold=True
		p.add_run(item['contenido'])
		p.add_run("\n")


doc.save("diccionario.docx")
